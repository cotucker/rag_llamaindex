import fitz
import io
import re
import os
import pymupdf
import markdown
import docx
import json
import pandas as pd
from PIL import Image
from src.image_captioning import caption_image, caption_image_groq
from llama_index.core import Document
from bs4 import BeautifulSoup

ACCESS_CONTROL_CONFIG = {}

def set_access_control_config(config: dict):
    global ACCESS_CONTROL_CONFIG
    ACCESS_CONTROL_CONFIG = config

def get_images_description(page) -> str:
    image_list = page.get_images(full=True)
    descriptions = []

    for image_index, img in enumerate(image_list, start=1):
        xref = img[0]
        base_image = page.parent.extract_image(xref)
        image_bytes = base_image["image"]
        caption = caption_image_groq(image_bytes)
        if not caption:
            continue
        desc_str = (
            f"Image: {caption.image_name} "
            f"Type: {caption.image_type} "
            f"Description: {caption.image_description}"
        )
        descriptions.append(desc_str)

    if not descriptions:
        return ""

    result = '\n'.join(descriptions)
    return result

def clean_text(text: str) -> str:
    text = re.sub(r'[^\w\s\.]', '', text)
    text = text.replace('\n', ' ')
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def clean_markdown(md_text: str) -> str:
    html = markdown.markdown(md_text)
    soup = BeautifulSoup(html, "html.parser")
    text = soup.get_text(separator=' ')
    return text.strip()

def get_document_from_pdf(path_to_pdf: str) -> Document:
    doc = pymupdf.open(path_to_pdf)
    text = ' '.join([page.get_text() + ' ' +  get_images_description(page) for page in doc])
    text = clean_text(text)
    file_name = os.path.basename(path_to_pdf)
    return Document(
        text=text,
        metadata={
            "file_path": path_to_pdf,
            "file_name": os.path.basename(path_to_pdf),
            "access_level": ACCESS_CONTROL_CONFIG.get(file_name, "private")
        }
    )

def get_document_from_txt(path_to_txt: str) -> Document:
    with open(path_to_txt, "r", encoding="utf-8") as f:
        text = f.read()
    text = clean_text(text)
    file_name = os.path.basename(path_to_txt)
    return Document(
        text=text,
        metadata={
            "file_path": path_to_txt,
            "file_name": file_name,
            "access_level": ACCESS_CONTROL_CONFIG.get(file_name, "private")
        }
    )

def get_document_from_md(path_to_md: str) -> Document:
    with open(path_to_md, "r", encoding="utf-8") as f:
        text = f.read()

    cleaned_markdown = clean_markdown(text)
    cleaned_text = clean_text(cleaned_markdown)
    file_name = os.path.basename(path_to_md)
    return Document(
        text=cleaned_text,
        metadata={
            "file_path": path_to_md,
            "file_name": file_name,
            "access_level": ACCESS_CONTROL_CONFIG.get(file_name, "private")
        }
    )

def get_document_from_docx(path_to_docx: str) -> Document:
    doc = docx.Document(path_to_docx)
    text_content = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_content.append(para.text)

    if doc.tables:
        text_content.append("\n--- TABLES DATA ---\n")
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                text_content.append(" | ".join(row_text))
            text_content.append("")

    full_text = "\n".join(text_content)
    file_name = os.path.basename(path_to_docx)
    cleaned_text = clean_text(full_text)
    return Document(
        text=cleaned_text,
        metadata={
            "file_path": path_to_docx,
            "file_name": file_name,
            "access_level": ACCESS_CONTROL_CONFIG.get(file_name, "private")
        }
    )

def get_document_from_csv(path_to_csv: str) -> Document:
    df = pd.read_csv(path_to_csv)
    text = df.to_string()
    text = clean_text(text)
    file_name = os.path.basename(path_to_csv)
    return Document(
        text=text,
        metadata={
            "file_path": path_to_csv,
            "file_name": file_name,
            "access_level": ACCESS_CONTROL_CONFIG.get(file_name, "private")
        }
    )

def get_document_from_xlsx(path_to_xlsx: str) -> Document:
    df = pd.read_excel(path_to_xlsx)
    text = df.to_string()
    text = clean_text(text)
    file_name = os.path.basename(path_to_xlsx)
    return Document(
        text=text,
        metadata={
            "file_path": path_to_xlsx,
            "file_name": file_name,
            "access_level": ACCESS_CONTROL_CONFIG.get(file_name, "private")
        }
    )

def get_document_from_image(path_to_image: str) -> Document | None:
    with open(path_to_image, "rb") as f:
        image_bytes = f.read()
    caption = caption_image(image_bytes)
    if not caption:
        return None
    text = (
        f"Image: {caption.image_name} "
        f"Type: {caption.image_type} "
        f"Description: {caption.image_description}"
    )
    file_name = os.path.basename(path_to_image)
    return Document(
        text=text,
        metadata={
            "file_path": path_to_image,
            "file_name": file_name,
            "access_level": ACCESS_CONTROL_CONFIG.get(file_name, "private")
        }
    )
